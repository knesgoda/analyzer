import streamlit as st
import json
import io
from typing import List
from docx import Document
from docx.shared import Pt
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

# --- CONFIGURATION ---
st.set_page_config(page_title="OutPaged Scene Generator", layout="wide")

# --- 1. CORE DATA STRUCTURES (From Analyzer Agent) ---
# We redefine these briefly here to ensure the app is self-contained
SKYBOX_NEGATIVE_PROMPT = (
    "people, person, faces, crowds, animals, text, letters, signage, "
    "watermark, logo, UI, placeable props, furniture, vehicles, "
    "modern objects, anachronistic items, blurry"
)

class Character:
    def __init__(self, name, role, visual_description):
        self.name = name
        self.role = role
        self.visual_description = visual_description

class Skybox:
    def __init__(self, visual_prompt, environment_type, negative_prompt=SKYBOX_NEGATIVE_PROMPT):
        self.visual_prompt = visual_prompt
        self.environment_type = environment_type
        self.negative_prompt = negative_prompt

class Scene:
    def __init__(self, location, chapter_beat, trigger_sentence, characters, skybox_environment):
        self.location = location
        self.chapter_beat = chapter_beat
        self.trigger_sentence = trigger_sentence
        self.characters = characters # List of Character objs
        self.skybox_environment = skybox_environment # Skybox obj

# --- 2. EPUB PARSING LOGIC ---
def extract_text_from_epub(uploaded_file):
    """
    Parses EPUB file and returns a list of chapter texts.
    Real implementation would need robust filtering for TOC/Copyright pages.
    """
    book = epub.read_epub(uploaded_file)
    chapters = []
    
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text(separator='\n').strip()
            # Basic filter to ignore tiny chapters (often TOC or spacers)
            if len(text) > 500: 
                chapters.append(text)
    
    return chapters

# --- 3. AI MOCK AGENT (The Brain Replacement) ---
def mock_analyze_chapter(text_chunk, chapter_num):
    """
    Simulates the AI's response to demonstrate the file generation logic.
    """
    # Create fake scenes based on the provided examples
    scenes = []
    
    # Scene A
    s1 = Scene(
        location="Mock Location A",
        chapter_beat=f"Chapter {chapter_num} - Part 1 - The Beginning",
        trigger_sentence=f"This is the trigger sentence found in chapter {chapter_num}, signaling the start.",
        characters=[
            Character("Main Hero", "Main", "A tall hero wearing a red cape. Full-body cutout, 4k."),
            Character("Sidekick", "Secondary", "A small sidekick with glasses. Full-body cutout, 4k.")
        ],
        skybox_environment=Skybox(
            visual_prompt=f"ground view center eye level, outdoors Mock City, sunny day, 3D watercolor",
            environment_type="Outdoors"
        )
    )
    scenes.append(s1)

    # Scene B
    s2 = Scene(
        location="Mock Location B",
        chapter_beat=f"Chapter {chapter_num} - Part 2 - The Climax",
        trigger_sentence=f"Suddenly, a loud noise echoed through the hall in chapter {chapter_num}.",
        characters=[
            Character("Main Hero", "Main", "A tall hero wearing a red cape. Full-body cutout, 4k."),
            Character("Villain", "Secondary", "A dark shadowy figure. Full-body cutout, 4k.")
        ],
        skybox_environment=Skybox(
            visual_prompt=f"ground view center eye level, indoors Dark Castle, candlelit, 3D watercolor",
            environment_type="Indoors"
        )
    )
    scenes.append(s2)
    
    return scenes

# --- 4. DOCUMENT GENERATION LOGIC ---
def generate_documents(book_title, all_scenes):
    """
    The engine that enforces the file naming and formatting rules.
    """
    
    # -- SETUP DOCS --
    doc1 = Document() # Triggers
    doc2 = Document() # Backgrounds
    doc3 = Document() # Characters
    
    # Headers
    doc1.add_heading(f"{book_title} | AR Scene Pack | Outputs v1", 0)
    doc1.add_paragraph("Document 1: Trigger Sentences Per Scene")
    
    doc2.add_heading(f"{book_title} | AR Scene Pack | Outputs v1", 0)
    doc2.add_paragraph("Document 2: Skybox Background Prompts (Blockade Labs Standard)")
    
    doc3.add_heading(f"{book_title} | AR Scene Pack | Outputs v1", 0)
    doc3.add_paragraph("Document 3: Character Descriptions (No Cues)")

    # -- ITERATE SCENES --
    global_idx = 1
    
    for scene in all_scenes:
        scene_id = f"ch{global_idx:02}" # The Logic: ch01, ch02, ch03...
        
        # --- DOC 1: TRIGGERS ---
        doc1.add_heading(f"Scene {global_idx:02} - {scene.location} - {scene.chapter_beat}", level=2)
        p = doc1.add_paragraph()
        p.add_run("Trigger sentence: ").bold = True
        p.add_run(scene.trigger_sentence)
        p = doc1.add_paragraph()
        p.add_run("Page number: ").bold = True
        p.add_run("N/A") # Default as per instructions
        
        # --- DOC 2: BACKGROUNDS ---
        doc2.add_heading(f"Scene {global_idx:02} - {scene.location} - {scene.chapter_beat}", level=2)
        doc2.add_paragraph(f"Background file name:  {scene_id}bg01")
        
        p = doc2.add_paragraph()
        p.add_run("Skybox prompt:  ").bold = True
        p.add_run(scene.skybox_environment.visual_prompt)
        
        p = doc2.add_paragraph()
        p.add_run("Negative prompt:  ").bold = True
        p.add_run(scene.skybox_environment.negative_prompt)
        
        doc2.add_paragraph() # Spacer
        
        # --- DOC 3: CHARACTERS ---
        doc3.add_heading(f"Scene {global_idx:02} - {scene.location} - {scene.chapter_beat}", level=2)
        
        # Main Character Logic (mc01)
        # We find the main character (role="Main") or default to first
        main_chars = [c for c in scene.characters if c.role == "Main"]
        other_chars = [c for c in scene.characters if c.role != "Main"]
        
        # If no main defined, take the first one
        if not main_chars and scene.characters:
            main_chars = [scene.characters[0]]
            other_chars = scene.characters[1:]
            
        # Write Main Character
        if main_chars:
            mc = main_chars[0]
            doc3.add_paragraph(f"{mc.name}:")
            doc3.add_paragraph(f"- ref: {scene_id}mc01") # Naming Rule: ch##mc01
            doc3.add_paragraph(f"prompt: \"{mc.visual_description}\"")
            
        # Write Secondary Characters
        sc_idx = 1
        for sc in other_chars:
            doc3.add_paragraph(f"{sc.name}:")
            doc3.add_paragraph(f"- ref: {scene_id}sc{sc_idx:02}") # Naming Rule: ch##sc##
            doc3.add_paragraph(f"prompt: \"{sc.visual_description}\"")
            sc_idx += 1
            
        global_idx += 1
        
    # -- SAVE TO BYTES --
    def save_to_stream(doc):
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    return save_to_stream(doc1), save_to_stream(doc2), save_to_stream(doc3)

# --- 5. UI LAYOUT ---
st.title("OutPaged Scene Pack Generator")
st.markdown("Automated EPUB ingestion and Scene Pack generation.")

with st.sidebar:
    st.header("Configuration")
    book_title = st.text_input("Book Title", "Winnie-the-Pooh")
    use_mock = st.checkbox("Use Mock Mode (No AI Cost)", value=True)
    api_key = st.text_input("OpenAI API Key", type="password", disabled=use_mock)
    st.info("Upload an EPUB to begin analysis.")

uploaded_file = st.file_uploader("Upload EPUB", type=["epub"])

if uploaded_file and st.button("Generate Scene Pack"):
    with st.spinner("Parsing EPUB..."):
        # 1. Parse
        chapters = extract_text_from_epub(uploaded_file)
        st.success(f"Found {len(chapters)} processable chapters.")
        
        # 2. Analyze (Loop)
        all_scenes = []
        progress_bar = st.progress(0)
        
        for i, chapter_text in enumerate(chapters):
            # Update UI
            progress_bar.progress((i + 1) / len(chapters))
            
            # Call Agent (Mock or Real)
            if use_mock:
                chapter_scenes = mock_analyze_chapter(chapter_text, i+1)
            else:
                # HERE IS WHERE WE WOULD CALL THE REAL ANALYZER AGENT
                # chapter_scenes = analyzer.analyze_chapter_content(client, chapter_text, book_title)
                chapter_scenes = mock_analyze_chapter(chapter_text, i+1) # Fallback for now
            
            all_scenes.extend(chapter_scenes)
            
        st.success(f"Analysis Complete! Generated {len(all_scenes)} scenes.")
        
        # 3. Generate Docs
        d1, d2, d3 = generate_documents(book_title, all_scenes)
        
        # 4. Download Columns
        c1, c2, c3 = st.columns(3)
        
        file_prefix = book_title.replace(" ", "_")
        
        with c1:
            st.download_button(
                label="📄 Download Trigger Doc",
                data=d1,
                file_name=f"{file_prefix}_Document_1_Triggers.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with c2:
            st.download_button(
                label="🖼️ Download Skybox Doc",
                data=d2,
                file_name=f"{file_prefix}_Document_2_Backgrounds.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with c3:
            st.download_button(
                label="👤 Download Character Doc",
                data=d3,
                file_name=f"{file_prefix}_Document_3_Characters.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
